"""DOCX parsing utilities."""

from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path
from typing import Any
from xml.etree import ElementTree
from zipfile import ZipFile

from docx import Document


@dataclass(slots=True)
class DocxParseResult:
    page_count: int
    tables: list[dict[str, Any]]
    sections: list[dict[str, Any]]
    full_text: str


def parse_docx(path: Path) -> DocxParseResult:
    document = Document(path)

    tables = _extract_tables(path, document)
    sections = _extract_sections(document)
    text_content = "\n".join(paragraph.text for paragraph in document.paragraphs if paragraph.text)
    page_count = _read_page_count(path)

    return DocxParseResult(page_count=page_count, tables=tables, sections=sections, full_text=text_content)


def _extract_tables(path: Path, document: Document) -> list[dict[str, Any]]:
    summaries: list[dict[str, Any]] = []
    for index, table in enumerate(document.tables, start=1):
        row_count = len(table.rows)
        column_count = len(table.columns) if table.columns else 0
        structured_rows = []
        headers = [cell.text.strip() if cell.text else f"col_{i}" for i, cell in enumerate(table.rows[0].cells)] if row_count > 0 else []

        for row in table.rows[1:]:
            row_map = {}
            for idx, cell in enumerate(row.cells):
                header = headers[idx] if idx < len(headers) else f"col_{idx}"
                row_map[header] = cell.text.strip() if cell.text else ""
            structured_rows.append(row_map)

        summaries.append(
            {
                "table_id": f"{path.stem}-table-{index}",
                "caption": None,
                "page_range": None,
                "row_count": len(structured_rows),
                "column_count": len(headers) if headers else column_count,
                "rows": structured_rows,
            }
        )
    return summaries


def _extract_sections(document: Document) -> list[dict[str, Any]]:
    sections: list[dict[str, Any]] = []
    seen_titles: set[str] = set()

    for paragraph in document.paragraphs:
        style_name = (paragraph.style.name if paragraph.style else "").lower()
        text = paragraph.text.strip()
        if not text:
            continue

        if style_name.startswith("heading"):
            level = _parse_heading_level(style_name)
        elif text.isupper() and len(text) <= 100:
            level = 1
        else:
            continue

        title = text if text not in seen_titles else f"{text} ({len(seen_titles)})"
        seen_titles.add(title)
        sections.append({"level": level, "title": title})

    return sections


def _parse_heading_level(style_name: str) -> int:
    parts = style_name.split()
    for part in parts:
        if part.isdigit():
            return int(part)
    return 1


def _read_page_count(path: Path) -> int:
    try:
        with ZipFile(path) as archive:
            with archive.open("docProps/app.xml") as properties:
                tree = ElementTree.parse(properties)
                namespace = {"ep": "http://schemas.openxmlformats.org/officeDocument/2006/extended-properties"}
                pages_text = tree.findtext("ep:Pages", namespaces=namespace)
                if pages_text and pages_text.isdigit():
                    return int(pages_text)
    except KeyError:
        return 0
    except Exception:
        return 0
    return 0
